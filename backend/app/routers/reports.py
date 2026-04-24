import os
import io
import json
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.datafile import DataFile
from app.routers.data import get_current_user
from app.models.user import User

router = APIRouter(prefix="/reports", tags=["Reports"])


def create_bar_chart(df: pd.DataFrame, numeric_cols: list) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 4))
    fig.patch.set_facecolor('#1e1e2e')
    ax.set_facecolor('#1e1e2e')

    x = range(len(df))
    colors = ['#6c63ff', '#10b981', '#f59e0b', '#ef4444']

    col = numeric_cols[0]
    bars = ax.bar(x, df[col], color=colors[0], alpha=0.85, width=0.6)

    ax.set_xlabel(df.columns[0], color='#94a3b8', fontsize=10)
    ax.set_ylabel(col, color='#94a3b8', fontsize=10)
    ax.set_title(f'Bar Chart — {col}', color='#e2e8f0', fontsize=12, fontweight='bold')

    if len(df) <= 20:
        ax.set_xticks(list(x))
        ax.set_xticklabels([str(v)[:10] for v in df.iloc[:, 0]], rotation=45, ha='right', color='#94a3b8', fontsize=8)
    ax.tick_params(colors='#94a3b8')
    ax.spines['bottom'].set_color('#334155')
    ax.spines['left'].set_color('#334155')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.grid(True, color='#2a2a3e', linestyle='--')
    ax.set_axisbelow(True)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#1e1e2e')
    plt.close()
    buf.seek(0)
    return buf


def create_line_chart(df: pd.DataFrame, numeric_cols: list) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 4))
    fig.patch.set_facecolor('#1e1e2e')
    ax.set_facecolor('#1e1e2e')

    colors = ['#6c63ff', '#10b981', '#f59e0b', '#ef4444']
    x = range(len(df))

    for i, col in enumerate(numeric_cols[:2]):
        color = colors[i % len(colors)]
        ax.plot(list(x), df[col].tolist(), color=color, linewidth=2.5, marker='o',
                markersize=4, label=col, alpha=0.9)

    ax.set_xlabel(df.columns[0], color='#94a3b8', fontsize=10)
    ax.set_ylabel('Values', color='#94a3b8', fontsize=10)
    ax.set_title('Line Chart — Trends Over Time', color='#e2e8f0', fontsize=12, fontweight='bold')

    if len(df) <= 20:
        ax.set_xticks(list(x))
        ax.set_xticklabels([str(v)[:10] for v in df.iloc[:, 0]], rotation=45, ha='right', color='#94a3b8', fontsize=8)
    ax.tick_params(colors='#94a3b8')
    ax.spines['bottom'].set_color('#334155')
    ax.spines['left'].set_color('#334155')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.grid(True, color='#2a2a3e', linestyle='--')
    ax.set_axisbelow(True)

    if len(numeric_cols) >= 2:
        legend = ax.legend(facecolor='#252535', labelcolor='#e2e8f0', edgecolor='#334155')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#1e1e2e')
    plt.close()
    buf.seek(0)
    return buf


def generate_pdf_report(db_file, df: pd.DataFrame, current_user: User) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    output_path = f"/app/uploads/report_{db_file.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    ACCENT = colors.HexColor('#6C63FF')
    DARK = colors.HexColor('#1e1e2e')
    LIGHT_BG = colors.HexColor('#f8f9fa')
    BORDER = colors.HexColor('#e2e8f0')

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'],
        fontSize=24, textColor=ACCENT, spaceAfter=6,
        alignment=TA_CENTER, fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'],
        fontSize=10, textColor=colors.HexColor('#94a3b8'),
        alignment=TA_CENTER, spaceAfter=20)
    section_style = ParagraphStyle('Section', parent=styles['Heading1'],
        fontSize=14, textColor=ACCENT, spaceBefore=16, spaceAfter=8,
        fontName='Helvetica-Bold')
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
        fontSize=10, textColor=colors.HexColor('#334155'), spaceAfter=4)

    story = []

    # Titre
    story.append(Paragraph("Data Analysis Report", title_style))
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} by {current_user.username}",
        subtitle_style
    ))
    story.append(Table([['']], colWidths=[17*cm],
        style=TableStyle([('LINEABOVE', (0,0), (-1,0), 1, ACCENT)])))
    story.append(Spacer(1, 12))

    # Infos fichier
    story.append(Paragraph("File Information", section_style))
    info_data = [
        ['File', db_file.original_filename],
        ['Total Rows', str(len(df))],
        ['Total Columns', str(len(df.columns))],
        ['File Size', f"{db_file.file_size} KB"],
        ['Uploaded', db_file.uploaded_at.strftime('%Y-%m-%d %H:%M')],
    ]
    info_table = Table(info_data, colWidths=[5*cm, 12*cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), LIGHT_BG),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 0.5, BORDER),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.white, LIGHT_BG]),
        ('PADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 16))

    # Statistiques numériques
    numeric_cols = df.select_dtypes(include='number').columns.tolist()
    if numeric_cols:
        story.append(Paragraph("Numeric Statistics", section_style))
        stats_df = df[numeric_cols].describe().round(2)
        stats_header = ['Stat'] + numeric_cols
        stats_data = [stats_header]
        for stat_name in stats_df.index:
            row = [stat_name] + [str(stats_df.loc[stat_name, col]) for col in numeric_cols]
            stats_data.append(row)
        col_width = 17*cm / len(stats_header)
        stats_table = Table(stats_data, colWidths=[col_width]*len(stats_header))
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), ACCENT),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, BORDER),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, LIGHT_BG]),
            ('PADDING', (0,0), (-1,-1), 6),
            ('ALIGN', (1,1), (-1,-1), 'CENTER'),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 20))

        # --- GRAPHIQUES ---
        story.append(Paragraph("Data Visualizations", section_style))
        story.append(Paragraph(
            "The following charts are automatically generated from the numeric columns in your dataset.",
            body_style
        ))
        story.append(Spacer(1, 8))

        # Graphique en barres
        story.append(Paragraph("Bar Chart", ParagraphStyle('ChartTitle',
            parent=styles['Normal'], fontSize=11, fontName='Helvetica-Bold',
            textColor=colors.HexColor('#334155'), spaceAfter=6)))
        bar_buf = create_bar_chart(df, numeric_cols)
        bar_img = Image(bar_buf, width=17*cm, height=7*cm)
        story.append(bar_img)
        story.append(Spacer(1, 16))

        # Graphique en courbes
        story.append(Paragraph("Line Chart", ParagraphStyle('ChartTitle',
            parent=styles['Normal'], fontSize=11, fontName='Helvetica-Bold',
            textColor=colors.HexColor('#334155'), spaceAfter=6)))
        line_buf = create_line_chart(df, numeric_cols)
        line_img = Image(line_buf, width=17*cm, height=7*cm)
        story.append(line_img)
        story.append(Spacer(1, 20))

    # Tableau de données
    story.append(Paragraph("Data Preview (first 50 rows)", section_style))
    story.append(Paragraph(f"Showing {min(50, len(df))} of {len(df)} rows.", body_style))

    preview_df = df.head(50).fillna('')
    table_data = [list(df.columns)]
    for _, row in preview_df.iterrows():
        table_data.append([str(v)[:20] for v in row.values])

    col_w = 17*cm / len(df.columns)
    data_table = Table(table_data, colWidths=[col_w]*len(df.columns))
    data_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), ACCENT),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('GRID', (0,0), (-1,-1), 0.5, BORDER),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, LIGHT_BG]),
        ('PADDING', (0,0), (-1,-1), 5),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
    ]))
    story.append(data_table)

    # Footer
    story.append(Spacer(1, 20))
    story.append(Table([['']], colWidths=[17*cm],
        style=TableStyle([('LINEABOVE', (0,0), (-1,0), 0.5, BORDER)])))
    story.append(Paragraph(
        "Generated by DataDash — Industrial Data Dashboard",
        subtitle_style
    ))

    doc.build(story)
    return output_path


def generate_word_report(db_file, df: pd.DataFrame, current_user: User) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    header = doc.add_heading('', 0)
    run = header.add_run('Data Analysis Report')
    run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
    run.font.size = Pt(24)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('─' * 60)

    meta = doc.add_paragraph()
    meta.add_run('File: ').bold = True
    meta.add_run(db_file.original_filename)
    meta.add_run('\n')
    meta.add_run('Generated: ').bold = True
    meta.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    meta.add_run('\n')
    meta.add_run('Generated by: ').bold = True
    meta.add_run(current_user.username)
    doc.add_paragraph()

    doc.add_heading('Summary Statistics', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for metric, value in [
        ('Total Rows', str(len(df))),
        ('Total Columns', str(len(df.columns))),
        ('File Size', f"{db_file.file_size} KB"),
        ('Columns', ', '.join(list(df.columns))),
    ]:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph()

    numeric_cols = df.select_dtypes(include='number').columns.tolist()
    if numeric_cols:
        doc.add_heading('Numeric Statistics', level=1)
        stats_df = df[numeric_cols].describe().round(2)
        num_table = doc.add_table(rows=1, cols=len(numeric_cols) + 1)
        num_table.style = 'Table Grid'
        header_row = num_table.rows[0].cells
        header_row[0].text = 'Stat'
        for i, col in enumerate(numeric_cols):
            header_row[i+1].text = col
            for p in header_row[i+1].paragraphs:
                for r in p.runs:
                    r.bold = True
        for stat_name in stats_df.index:
            row_cells = num_table.add_row().cells
            row_cells[0].text = stat_name
            for i, col in enumerate(numeric_cols):
                row_cells[i+1].text = str(stats_df.loc[stat_name, col])
        doc.add_paragraph()

        # Graphiques dans Word
        doc.add_heading('Data Visualizations', level=1)

        doc.add_heading('Bar Chart', level=2)
        bar_buf = create_bar_chart(df, numeric_cols)
        doc.add_picture(bar_buf, width=Inches(6))
        doc.add_paragraph()

        doc.add_heading('Line Chart', level=2)
        line_buf = create_line_chart(df, numeric_cols)
        doc.add_picture(line_buf, width=Inches(6))
        doc.add_paragraph()

    doc.add_heading('Data Table', level=1)
    doc.add_paragraph(f'Showing first 50 rows of {len(df)} total rows.')
    preview_df = df.head(50)
    data_table = doc.add_table(rows=1, cols=len(df.columns))
    data_table.style = 'Table Grid'
    header_row = data_table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_row[i].text = str(col)
        for p in header_row[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for _, row in preview_df.iterrows():
        row_cells = data_table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by DataDash — Industrial Data Dashboard')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run.font.size = Pt(9)

    output_path = f"/app/uploads/report_{db_file.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(output_path)
    return output_path


@router.post("/generate/{file_id}")
async def generate_report(
    file_id: int,
    format: str = "pdf",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_file = db.query(DataFile).filter(
        DataFile.id == file_id,
        DataFile.user_id == current_user.id
    ).first()

    if not db_file:
        raise HTTPException(status_code=404, detail="File not found")

    try:
        if db_file.original_filename.endswith('.csv'):
            df = pd.read_csv(db_file.file_path)
        else:
            df = pd.read_excel(db_file.file_path)
        df = df.fillna('')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not read file: {str(e)}")

    try:
        if format == "word":
            output_path = generate_word_report(db_file, df, current_user)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"report_{db_file.original_filename.replace('.csv','').replace('.xlsx','')}.docx"
        else:
            output_path = generate_pdf_report(db_file, df, current_user)
            media_type = "application/pdf"
            filename = f"report_{db_file.original_filename.replace('.csv','').replace('.xlsx','')}.pdf"

        return FileResponse(
            path=output_path,
            media_type=media_type,
            filename=filename,
            background=None
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")